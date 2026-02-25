"""
retirement_core.py
Core math, dataclasses, and plotting for the retirement planner.
No Streamlit dependencies here.
"""
import copy
from dataclasses import dataclass, asdict
from typing import Dict, List, Tuple
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches


@dataclass
class Inputs:
    current_age: int = 40
    pension_access_age: int = 57
    state_pension_age: int = 67
    life_expectancy: int = 85
    annual_spending_today: float = 20000.0
    inflation_rate: float = 0.025
    nominal_return: float = 0.05
    include_state_pension: bool = True
    state_pension_annual_today: float = 11502.0
    sizing_method: str = "SWR"  # "SWR" or "PV"
    swr: float = 0.04
    current_isa: float = 0.0
    current_pension: float = 0.0
    # Desired pension balance (today's £) at life_expectancy
    desired_pension_at_end: float = 0.0

    # --- NEW: Scenario triggers ---
    adj_swr: float = 0.035
    adj_inflation: float = 0.005
    adj_return: float = 0.005
    # --- NEW: Other Income ---
    other_income: float = 0.0
    other_income_years: int = 0


@dataclass
class Results:
    real_return: float
    bridge_years: int
    isa_needed_today: float
    pension_needed_today: float
    total_needed_today: float
    pension_needed_at_access: float
    mathematical_min_total_today: float  # Only used for comparison if method="SWR"


def _safe_pv_annuity(payment: float, r: float, n: int) -> float:
    if n <= 0:
        return 0.0
    if abs(r) < 1e-9:
        return payment * n
    return payment * (1 - (1 + r) ** (-n)) / r


def _real_return(nominal_return: float, inflation_rate: float) -> float:
    return (1 + nominal_return) / (1 + inflation_rate) - 1


def size_capital(inputs: Inputs) -> Tuple[Results, pd.DataFrame, pd.DataFrame]:
    """
    Compute required capital and produce deterministic projections.
    Returns:
        results: summary numbers
        projection: year-by-year cashflow projection
        meta: one-row dataframe with inputs + results
    """
    # Validation
    if inputs.pension_access_age < inputs.current_age:
        raise ValueError("Pension access age must be >= current age")
    if inputs.life_expectancy < inputs.current_age:
        raise ValueError("Life expectancy must be >= current age")
    if inputs.annual_spending_today < 0:
        raise ValueError("Annual spending must be >= 0")
    if inputs.inflation_rate <= -0.99:
        raise ValueError("Inflation rate too low")
    if inputs.nominal_return <= -0.99:
        raise ValueError("Nominal return too low")
    if inputs.sizing_method not in {"PV", "SWR"}:
        raise ValueError("sizing_method must be 'PV' or 'SWR'")
    if inputs.sizing_method == "SWR" and not (0 < inputs.swr < 0.2):
        raise ValueError("SWR must be between 0 and 20%")
    
    rr = _real_return(inputs.nominal_return, inputs.inflation_rate)
    bridge_years = inputs.pension_access_age - inputs.current_age
    
    # ISA bridge sizing
    isa_needed_today = _safe_pv_annuity(
        inputs.annual_spending_today, rr, bridge_years) * (1 + rr)
        
    # --- NEW: Other Income Offset (ISA Phase) ---
    isa_oi_yrs = min(bridge_years, inputs.other_income_years)
    pv_oi_isa = _safe_pv_annuity(inputs.other_income, inputs.nominal_return, isa_oi_yrs) * (1 + inputs.nominal_return)
    isa_needed_today = max(0.0, isa_needed_today - pv_oi_isa)
    
    # Post-access phases
    years_access_to_state = max(
        0, inputs.state_pension_age - inputs.pension_access_age)
    years_state_to_end = max(
        0, inputs.life_expectancy - inputs.state_pension_age + 1)
    
    pv_pre_state_at_access = _safe_pv_annuity(
        inputs.annual_spending_today, rr, years_access_to_state
    ) * (1 + rr)    
    
    if inputs.include_state_pension:
        gap_today = max(
            0.0, inputs.annual_spending_today - inputs.state_pension_annual_today
        )
    else:
        gap_today = inputs.annual_spending_today
        
    pv_gap_at_state_age = _safe_pv_annuity(gap_today, rr, years_state_to_end) * (1 + rr)
    
    pv_gap_at_access = (
        pv_gap_at_state_age / ((1 + rr) ** years_access_to_state)
        if years_access_to_state > 0
        else pv_gap_at_state_age
    )
    
    # Desired end-of-life pension balance, discounted to access age using real return
    years_access_to_end = max(
        0, inputs.life_expectancy - inputs.pension_access_age + 1)
        
    pv_end_at_access = (
        inputs.desired_pension_at_end / ((1 + rr) ** years_access_to_end)
        if years_access_to_end > 0
        else inputs.desired_pension_at_end
    )
    
    pension_needed_at_access_pv = (
        pv_pre_state_at_access + pv_gap_at_access + pv_end_at_access
    )
    pension_needed_today_pv = (
        pension_needed_at_access_pv / ((1 + rr) ** bridge_years)
        if bridge_years > 0
        else pension_needed_at_access_pv
    )
    
    mathematical_min_total_today = isa_needed_today + pension_needed_today_pv

    # SWR sizing (add desired end-balance as extra lump sum at access)
    if inputs.sizing_method == "SWR":
        mult = 1.0 / inputs.swr
    
        if years_state_to_end > 0:
            # You reach State Pension age during the model horizon
            if years_access_to_state > 0:
                # Before State Pension: you need full spending; after: only the gap
                temp_shortfall = inputs.annual_spending_today - gap_today  # typically the State Pension amount (capped by spending)
                pv_temp = _safe_pv_annuity(temp_shortfall, rr, years_access_to_state) * (1 + rr)
                pension_needed_at_access_base = (gap_today * mult) + pv_temp
            else:
                # State Pension already active at pension access
                pension_needed_at_access_base = gap_today * mult
        else:
            # You never reach State Pension age (so it must not reduce required pot)
            pension_needed_at_access_base = inputs.annual_spending_today * mult
    
        pension_needed_at_access = pension_needed_at_access_base + pv_end_at_access
        pension_needed_today = (
            pension_needed_at_access / ((1 + rr) ** bridge_years)
            if bridge_years > 0
            else pension_needed_at_access
        )
    else:
        pension_needed_at_access = pension_needed_at_access_pv
        pension_needed_today = pension_needed_today_pv

    # --- NEW: Other Income Offset (Pension Phase) ---
    pens_oi_yrs = max(0, inputs.other_income_years - bridge_years)
    if pens_oi_yrs > 0:
        pv_oi_pens_at_access = _safe_pv_annuity(inputs.other_income, inputs.nominal_return, pens_oi_yrs) * (1 + inputs.nominal_return)
        pv_oi_pens_today = pv_oi_pens_at_access / ((1 + inputs.nominal_return) ** bridge_years) if bridge_years > 0 else pv_oi_pens_at_access
        
        pension_needed_at_access = max(0.0, pension_needed_at_access - pv_oi_pens_at_access)
        pension_needed_today = max(0.0, pension_needed_today - pv_oi_pens_today)

    total_needed_today = isa_needed_today + pension_needed_today
    results = Results(
        real_return=rr,
        bridge_years=bridge_years,
        isa_needed_today=isa_needed_today,
        pension_needed_today=pension_needed_today,
        total_needed_today=total_needed_today,
        pension_needed_at_access=pension_needed_at_access,
        mathematical_min_total_today=mathematical_min_total_today,
    )
    
    # Deterministic projection with realistic depletion:
    projection: List[Dict] = []
    bal_isa = isa_needed_today
    bal_pension = pension_needed_today

    has_actuals = inputs.current_isa > 0 or inputs.current_pension > 0
    act_isa = inputs.current_isa
    act_pension = inputs.current_pension

    for t in range(inputs.life_expectancy - inputs.current_age + 1):
        age = inputs.current_age + t
        infl_factor = (1 + inputs.inflation_rate) ** t
        spending_nominal = inputs.annual_spending_today * infl_factor

        state_pension_nominal = 0.0
        if inputs.include_state_pension and age >= inputs.state_pension_age:
            state_pension_nominal = inputs.state_pension_annual_today * infl_factor

        # --- NEW: Other Income (Nominal) ---
        other_income_nominal = inputs.other_income if t < inputs.other_income_years else 0.0

        withdrawal_needed = max(0.0, spending_nominal - state_pension_nominal - other_income_nominal)

        if age < inputs.pension_access_age:
            isa_withdrawal = max(0.0, spending_nominal - other_income_nominal)
            bal_isa -= isa_withdrawal
            bal_isa *= (1 + inputs.nominal_return)
            bal_pension *= (1 + inputs.nominal_return)
            source = "ISA"
            portfolio_withdrawal = isa_withdrawal
            
            if has_actuals:
                act_isa -= isa_withdrawal
                if act_isa > 0:
                    act_isa *= (1 + inputs.nominal_return)
                act_pension *= (1 + inputs.nominal_return)
        else:
            bal_isa *= (1 + inputs.nominal_return)
            bal_pension -= withdrawal_needed
            bal_pension *= (1 + inputs.nominal_return)
            source = "Pension"
            portfolio_withdrawal = withdrawal_needed

            if has_actuals:
                if act_isa > 0:
                    act_isa *= (1 + inputs.nominal_return)
                act_pension -= withdrawal_needed
                act_pension *= (1 + inputs.nominal_return)

        total_bal = bal_isa + bal_pension
        act_total = (act_isa + act_pension) if has_actuals else None

        projection.append(
            {
                "Age": age,
                "YearIndex": t,
                "Spending": spending_nominal,
                "StatePension": state_pension_nominal,
                "OtherIncome": other_income_nominal,
                "PortfolioWithdrawal": portfolio_withdrawal,
                "Source": source,
                "ISA_Balance": bal_isa,
                "Pension_Balance": bal_pension,
                "Total_Balance": total_bal,
                "Actual_Total_Balance": act_total,
                # --- NEW: Track separate actuals ---
                "Actual_ISA_Balance": act_isa if has_actuals else None,
                "Actual_Pension_Balance": act_pension if has_actuals else None,
            }
        )
        
    df = pd.DataFrame(projection)
    
    # Build meta df
    meta_dict = vars(inputs).copy()
    meta_dict.update(vars(results))
    meta_df = pd.DataFrame([meta_dict])
    return results, df, meta_df


def _fmt_k(val: float, _: int) -> str:
    """Formatter for matplotlib axes (e.g. 500000 -> 500k)"""
    return f"{val/1000:,.0f}k"

def plot_single_projection(
    inputs: Inputs, results: Results, df: pd.DataFrame, out_png: str
) -> None:
    """Single-scenario visualisation (PNG)."""
    fig, axes = plt.subplots(2, 1, figsize=(13, 9), sharex=True)

    # ----------------------------
    # (1) Balances
    # ----------------------------
    ax = axes[0]
    ax.plot(df["Age"], df["Total_Balance"], lw=2.6, label="Total (Target)", color="#2E86AB")
    ax.plot(df["Age"], df["ISA_Balance"], lw=1.8, label="ISA (Target)", color="#A23B72")
    ax.plot(df["Age"], df["Pension_Balance"], lw=1.8, label="Pension (Target)", color="#F18F01")

    # Optional actuals (kept from your newer build if present)
    if "Actual_Total_Balance" in df.columns and df["Actual_Total_Balance"].notna().any():
        ax.plot(df["Age"], df["Actual_Total_Balance"], lw=2.6, label="Actual Total", color="black", ls="--")
        if "Actual_ISA_Balance" in df.columns:
            ax.plot(df["Age"], df["Actual_ISA_Balance"], lw=1.8, label="Actual ISA", color="#A23B72", ls=":")
        if "Actual_Pension_Balance" in df.columns:
            ax.plot(df["Age"], df["Actual_Pension_Balance"], lw=1.8, label="Actual Pension", color="#F18F01", ls=":")

        # Highlight ISA Shortfall before pension access (as per old behaviour)
        if "Actual_ISA_Balance" in df.columns:
            bridge_mask = df["Age"] <= inputs.pension_access_age
            if (df.loc[bridge_mask, "Actual_ISA_Balance"] < 0).any():
                ax.fill_between(
                    df.loc[bridge_mask, "Age"],
                    0,
                    df.loc[bridge_mask, "Actual_ISA_Balance"],
                    where=df.loc[bridge_mask, "Actual_ISA_Balance"] < 0,
                    color="red",
                    alpha=0.25,
                    label="ISA Shortfall (Illiquid)",
                    interpolate=True,
                )

    ax.axhline(0, color="red", ls="--", lw=1, alpha=0.6)

    # REINSTATED: dotted phase markers on the balances chart (missing in your “new” screenshot)
    ax.axvline(inputs.pension_access_age, color="orange", ls=":", lw=2)
    if inputs.include_state_pension:
        ax.axvline(inputs.state_pension_age, color="green", ls=":", lw=2)

    ax.set_title("Portfolio balances over time", fontsize=13, fontweight="bold")
    ax.set_ylabel("Balance (£)", fontsize=11)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(_fmt_k))
    ax.grid(True, alpha=0.25)
    ax.legend(loc="upper left", bbox_to_anchor=(1.02, 1.0), borderaxespad=0.0, fontsize=9)

    # ----------------------------
    # (2) Income / spending
    # ----------------------------
    ax = axes[1]
    ax.plot(df["Age"], df["Spending"], lw=2.2, color="red", ls="--", label="Spending")

    ax.fill_between(
        df["Age"],
        0,
        df["StatePension"],
        alpha=0.45,
        color="#06A77D",
        label="State pension",
    )

    # NEW: Other income shading (blue), stacked above State Pension
    other = df["OtherIncome"] if "OtherIncome" in df.columns else 0.0
    ax.fill_between(
        df["Age"],
        df["StatePension"],
        df["StatePension"] + other,
        alpha=0.30,
        color="#2E86AB",
        label="Other income",
    )

    ax.fill_between(
        df["Age"],
        df["StatePension"] + other,
        df["StatePension"] + other + df["PortfolioWithdrawal"],
        alpha=0.45,
        color="#D4B483",
        label="Portfolio withdrawal",
    )

    ax.axvline(inputs.pension_access_age, color="orange", ls=":", lw=2, label="Pension access")
    if inputs.include_state_pension:
        ax.axvline(inputs.state_pension_age, color="green", ls=":", lw=2, label="State pension starts")

    ax.set_title("Annual spending and income sources", fontsize=13, fontweight="bold")
    ax.set_xlabel("Age (years)", fontsize=11)
    ax.set_ylabel("Annual amount (£)", fontsize=11)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(_fmt_k))
    ax.grid(True, alpha=0.25)
    ax.legend(loc="upper left", bbox_to_anchor=(1.02, 1.0), borderaxespad=0.0, fontsize=9)

    fig.suptitle(
        f"Retirement projection | sizing={inputs.sizing_method}"
        f" | return={inputs.nominal_return:.1%} | inflation={inputs.inflation_rate:.1%}"
        f" | spending=£{inputs.annual_spending_today:,.0f} (today)",
        y=0.98,
        fontsize=12,
    )
    fig.tight_layout(rect=[0, 0, 0.82, 0.96])
    fig.savefig(out_png, dpi=200, bbox_inches="tight")
    plt.close(fig)

def plot_scenario_set(base: Inputs, out_png: str) -> pd.DataFrame:
    """Scenario-comparison visualisation (PNG) + summary dataframe."""
    scenarios: List[Tuple[str, Inputs]] = []
    scenarios.append(("Base (SWR 4%)", base))
    scenarios.append(
        (
            f"Conservative (SWR {base.adj_swr*100:.1f}%)",
            Inputs(**{**asdict(base), "swr": base.adj_swr, "sizing_method": "SWR"}),
        )
    )
    scenarios.append(("PV (mathematical min)", Inputs(**{**asdict(base), "sizing_method": "PV"})))
    scenarios.append(
        (
            f"Inflation adjustment ({base.adj_inflation*100:+.2f}%)",
            Inputs(**{**asdict(base), "inflation_rate": base.inflation_rate + base.adj_inflation}),
        )
    )
    scenarios.append(
        (
            f"Returns adjustment ({base.adj_return*100:+.2f}%)",
            Inputs(**{**asdict(base), "nominal_return": base.nominal_return + base.adj_return}),
        )
    )

    summary_rows: List[Dict] = []
    scenario_results: List[Tuple[str, Inputs, Results, pd.DataFrame]] = []
    for label, inp in scenarios:
        res, df, _meta = size_capital(inp)
        summary_rows.append(
            {
                "Scenario": label,
                "TotalNeededToday": res.total_needed_today,
                "ISANeededToday": res.isa_needed_today,
                "PensionNeededToday": res.pension_needed_today,
            }
        )
        scenario_results.append((label, inp, res, df))

    summary = pd.DataFrame(summary_rows)

    fig, axes = plt.subplots(2, 2, figsize=(16, 11))

    # (1) Total balance over time
    ax = axes[0, 0]
    for label, _inp, _res, df_s in scenario_results:
        ax.plot(df_s["Age"], df_s["Total_Balance"], lw=2, label=label)
    ax.axhline(0, color="red", ls="--", lw=1, alpha=0.6)
    ax.set_title("Total balance over time", fontweight="bold")
    ax.set_xlabel("Age (years)")
    ax.set_ylabel("Balance (£)")
    ax.yaxis.set_major_formatter(plt.FuncFormatter(_fmt_k))
    ax.grid(True, alpha=0.25)
    ax.legend(fontsize=8, loc="upper left", bbox_to_anchor=(1.02, 1.0), borderaxespad=0.0)

    # (2) ISA vs pension needed today (bar chart)
    ax = axes[0, 1]
    x = np.arange(len(summary))
    w = 0.38
    ax.bar(x - w / 2, summary["ISANeededToday"], width=w, label="ISA", color="#A23B72")
    ax.bar(x + w / 2, summary["PensionNeededToday"], width=w, label="Pension", color="#F18F01")
    ax.set_title("Capital needed today", fontweight="bold")
    ax.set_xticks(x)
    ax.set_xticklabels(summary["Scenario"], rotation=30, ha="right", fontsize=8)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(_fmt_k))
    ax.grid(True, alpha=0.25, axis="y")
    ax.legend(loc="upper left", bbox_to_anchor=(1.02, 1.0), borderaxespad=0.0)

    # (3) Base scenario income sources (now includes Other Income)
    base_label, base_inp, _base_res, base_df = scenario_results[0]
    ax = axes[1, 0]
    ax.fill_between(base_df["Age"], 0, base_df["StatePension"], alpha=0.45, color="#06A77D", label="State pension")

    other = base_df["OtherIncome"] if "OtherIncome" in base_df.columns else 0.0
    ax.fill_between(
        base_df["Age"],
        base_df["StatePension"],
        base_df["StatePension"] + other,
        alpha=0.30,
        color="#2E86AB",
        label="Other income",
    )
    ax.fill_between(
        base_df["Age"],
        base_df["StatePension"] + other,
        base_df["StatePension"] + other + base_df["PortfolioWithdrawal"],
        alpha=0.45,
        color="#D4B483",
        label="Portfolio",
    )
    ax.plot(base_df["Age"], base_df["Spending"], color="red", ls="--", lw=2, label="Spending")
    ax.axvline(base_inp.pension_access_age, color="orange", ls=":", lw=2, label="Pension access")
    if base_inp.include_state_pension:
        ax.axvline(base_inp.state_pension_age, color="green", ls=":", lw=2, label="State pension")
    ax.set_title("Base scenario: spending and income", fontweight="bold")
    ax.set_xlabel("Age (years)")
    ax.set_ylabel("Annual amount (£)")
    ax.yaxis.set_major_formatter(plt.FuncFormatter(_fmt_k))
    ax.grid(True, alpha=0.25)
    ax.legend(fontsize=8, loc="upper left", bbox_to_anchor=(1.02, 1.0), borderaxespad=0.0)

    # (4) Summary table
    ax = axes[1, 1]
    ax.axis("off")
    tdf = summary.copy()
    for c in ["TotalNeededToday", "ISANeededToday", "PensionNeededToday"]:
        tdf[c] = tdf[c].map(lambda v: f"£{v:,.0f}")
    tdf.columns = ["Scenario", "Total\nNeeded Today", "ISA\nNeeded Today", "Pension\nNeeded Today"]
    tbl = ax.table(cellText=tdf.values.tolist(), colLabels=list(tdf.columns), cellLoc="center", loc="center")
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(9)
    tbl.scale(1.0, 1.9)
    tbl.auto_set_column_width(col=list(range(len(tdf.columns))))
    ax.set_title("Capital needed today", fontweight="bold")

    fig.suptitle("Retirement scenario comparison", fontsize=14, fontweight="bold")
    fig.tight_layout(rect=[0, 0, 0.82, 0.96])
    fig.savefig(out_png, dpi=200, bbox_inches="tight")
    plt.close(fig)

    return summary

def create_pension_report(
    out_docx,
    inputs: Inputs,
    results: Results,
    df: pd.DataFrame,
    chart_png: str,
    scenario_png: str,
):
    """Generate Word Document."""
    doc = Document()
    doc.add_heading("Retirement Plan – Output Report", 0)

    doc.add_heading("Target Capital Required Today", level=1)
    p = doc.add_paragraph()
    p.add_run(f"ISA / accessible needed: £{results.isa_needed_today:,.0f}\n").bold = True
    p.add_run(f"Pension needed: £{results.pension_needed_today:,.0f}\n").bold = True
    p.add_run(f"Total needed today: £{results.total_needed_today:,.0f}\n").bold = True

    doc.add_heading("Assumptions Summary", level=1)
    doc.add_paragraph(
        f"Ages: Current {inputs.current_age} | Access {inputs.pension_access_age} "
        f"| State {inputs.state_pension_age} | Life exp {inputs.life_expectancy}"
    )
    doc.add_paragraph(
        f"Spending: £{inputs.annual_spending_today:,.0f} (today's £) "
        f"| Inflation {inputs.inflation_rate*100:.1f}% "
        f"| Return {inputs.nominal_return*100:.1f}%"
    )
    if inputs.include_state_pension:
        doc.add_paragraph(f"State Pension included: £{inputs.state_pension_annual_today:,.0f} / yr (today's £)")
    if inputs.other_income > 0:
        doc.add_paragraph(f"Other Income: £{inputs.other_income:,.0f} for {inputs.other_income_years} years (nominal)")
    if inputs.desired_pension_at_end > 0:
        doc.add_paragraph(f"Longevity buffer (today's £) at life expectancy: £{inputs.desired_pension_at_end:,.0f}")
        
    doc.add_paragraph(f"Sizing methodology: {inputs.sizing_method}")

    doc.add_heading("Base Scenario Visualisation", level=1)
    doc.add_picture(chart_png, width=Inches(6.0))

    doc.add_heading("Scenario Sensitivities", level=1)
    doc.add_picture(scenario_png, width=Inches(6.0))

    doc.save(out_docx)
